"""
================================================================================
PROYECTO: Generador de Demanda de Amparo Indirecto contra la LMTR
VERSIÓN: 1.0
FECHA: Agosto 2026
LICENCIA: Creative Commons Atribución-CompartirIgual 4.0 Internacional (CC BY-SA 4.0)

DESCRIPCIÓN GENERAL:
Aplicación web construida con Streamlit para la automatización, personalización
y generación de demandas de amparo indirecto frente a la Ley en Materia de 
Telecomunicaciones y Radiodifusión (LMTR) y los Lineamientos de Identificación 
de Líneas Telefónicas Móviles. 

El sistema ejecuta el cómputo de plazos procesales del calendario judicial 2026,
la manipulación no destructiva de documentos OpenXML (.docx), el formateo 
tipográfico selectivo de rubros y pruebas, la preservación de notas al pie y la
sanitización de archivos para el Portal de Servicios en Línea del PJF.

CRÉDITOS Y DESARROLLO:
Diseñado y desarrollado bajo las instrucciones, dirección jurídica y supervisión 
estricta de Julio Amador, en colaboración y asistencia técnica de modelos de 
lenguaje (Gemini, Gemini Notebook y Grok).
================================================================================
"""

import io
import re
import unicodedata
from datetime import date, datetime, time, timedelta
from docx import Document
import streamlit as st

# Inicialización segura de la hora por defecto en sesión
if "hora_defecto" not in st.session_state:
    st.session_state.hora_defecto = datetime.now().time().replace(second=0, microsecond=0)

# ============================================================
# 1. CÓMPUTO DE PLAZOS Y FECHAS JUDICIALES
# ============================================================
def es_habil(fecha):
    """Determina si una fecha corresponde a un día hábil judicial en 2026."""
    if fecha.weekday() >= 5:  # Sábado o Domingo
        return False
    inhabiles = [
        date(2026, 9, 14), date(2026, 9, 15), date(2026, 9, 16),
        date(2026, 10, 12),
        date(2026, 11, 2), date(2026, 11, 16), date(2026, 11, 20),
    ] + [date(2026, 12, d) for d in range(16, 31)]
    return fecha not in inhabiles

def siguiente_habil(fecha):
    """Avanza secuencialmente hasta encontrar el siguiente día hábil."""
    while not es_habil(fecha):
        fecha += timedelta(days=1)
    return fecha

def obtener_fecha_base_calendario(numero_telefono):
    """Obtiene la fecha base oficial de desconexión según el último dígito."""
    digitos = [c for c in str(numero_telefono) if c.isdigit()]
    if not digitos:
        return date(2026, 8, 15)
    
    ultimo_digito = digitos[-1]
    mapa_fechas_base = {
        '0': date(2026, 8, 15),
        '1': date(2026, 8, 31),
        '2': date(2026, 9, 15),
        '3': date(2026, 9, 30),
        '4': date(2026, 10, 15),
        '5': date(2026, 10, 31),
        '6': date(2026, 11, 15),
        '7': date(2026, 11, 30),
        '8': date(2026, 12, 15),
        '9': date(2026, 12, 31),
    }
    return mapa_fechas_base.get(ultimo_digito, date(2026, 8, 15))

def calcular_plazos(fecha_conocimiento):
    """Calcula el cómputo de 15 días hábiles conforme a la Ley de Amparo."""
    fecha_surte = siguiente_habil(fecha_conocimiento + timedelta(days=1))
    fecha_empieza = siguiente_habil(fecha_surte + timedelta(days=1))
    
    fecha_final = fecha_empieza
    dias = 1
    while dias < 15:
        fecha_final += timedelta(days=1)
        if es_habil(fecha_final):
            dias += 1
    
    def fmt(d):
        meses = ["enero","febrero","marzo","abril","mayo","junio",
                 "julio","agosto","septiembre","octubre","noviembre","diciembre"]
        return f"{d.day} de {meses[d.month-1]} de {d.year}"
    
    return {
        "str_notif": fmt(fecha_conocimiento),
        "str_surte": fmt(fecha_surte),
        "str_empieza": fmt(fecha_empieza),
        "str_final": fmt(fecha_final),
        "fecha_final_obj": fecha_final
    }

# ============================================================
# 2. MOTOR DE REEMPLAZO Y FORMATEO TIPOGRÁFICO DE WORD
# ============================================================
def set_run_text_with_breaks(run, text):
    """Inserta saltos de línea estructurados (<w:br/>) en un run de Word."""
    lines = str(text).split('\n')
    run.text = lines[0]
    for line in lines[1:]:
        run.add_break()
        if line:
            run.add_text(line)

def obtener_todos_los_parrafos(doc):
    """Generador que itera sobre todos los párrafos de cuerpo, tablas y secciones."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                yield para
        if section.footer:
            for para in section.footer.paragraphs:
                yield para

def reemplazar_en_parrafo(para, placeholder, valor):
    """Reemplaza placeholders manejando fragmentación de runs en Word."""
    val_str = str(valor) if valor is not None else ""
    
    for run in para.runs:
        if placeholder in run.text:
            if '\n' in val_str:
                new_text = run.text.replace(placeholder, val_str)
                set_run_text_with_breaks(run, new_text)
            else:
                run.text = run.text.replace(placeholder, val_str)
            return

    while placeholder in para.text:
        full_text = "".join(r.text for r in para.runs)
        start_idx = full_text.find(placeholder)
        if start_idx == -1:
            break
        end_idx = start_idx + len(placeholder)
        
        current_pos = 0
        first_run_idx = None
        last_run_idx = None
        
        for idx, run in enumerate(para.runs):
            run_len = len(run.text)
            run_end = current_pos + run_len
            if first_run_idx is None and start_idx < run_end:
                first_run_idx = idx
            if last_run_idx is None and end_idx <= run_end:
                last_run_idx = idx
                break
            current_pos = run_end
            
        if first_run_idx is None or last_run_idx is None:
            break
            
        pos_in_first = sum(len(para.runs[i].text) for i in range(first_run_idx))
        rel_start = start_idx - pos_in_first
        
        pos_in_last = sum(len(para.runs[i].text) for i in range(last_run_idx))
        rel_end = end_idx - pos_in_last
        
        if first_run_idx == last_run_idx:
            run = para.runs[first_run_idx]
            new_text = run.text[:rel_start] + val_str + run.text[rel_end:]
            if '\n' in val_str:
                set_run_text_with_breaks(run, new_text)
            else:
                run.text = new_text
        else:
            first_run = para.runs[first_run_idx]
            last_run = para.runs[last_run_idx]
            
            new_first_text = first_run.text[:rel_start] + val_str
            if '\n' in val_str:
                set_run_text_with_breaks(first_run, new_first_text)
            else:
                first_run.text = new_first_text
                
            last_run.text = last_run.text[rel_end:]
            for i in range(first_run_idx + 1, last_run_idx):
                para.runs[i].text = ""

def reemplazar(doc, placeholder, valor):
    for para in obtener_todos_los_parrafos(doc):
        if placeholder in para.text:
            reemplazar_en_parrafo(para, placeholder, valor)

def reemplazar_frase(doc, frase_buscar, frase_reemplazo):
    reemplazar(doc, frase_buscar, frase_reemplazo)

def eliminar_seccion(doc, inicio, fin):
    """Elimina párrafos delimitados entre dos marcadores específicos."""
    eliminando = False
    to_remove = []
    for para in obtener_todos_los_parrafos(doc):
        texto = para.text
        if inicio in texto:
            eliminando = True
        if eliminando:
            to_remove.append(para)
            if fin in texto:
                break
    for para in to_remove:
        p = para._element
        parent = p.getparent()
        if parent is not None:
            if parent.tag.endswith('tc') and len(parent.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')) <= 1:
                para.text = ""
            else:
                parent.remove(p)

def eliminar_todas_las_secciones(doc, inicio, fin):
    while True:
        encontrado = False
        for para in obtener_todos_los_parrafos(doc):
            if inicio in para.text:
                encontrado = True
                break
        if not encontrado:
            break
        eliminar_seccion(doc, inicio, fin)

def borrar_frase_entre_marcadores(doc, inicio, fin):
    for para in obtener_todos_los_parrafos(doc):
        texto = para.text
        if inicio in texto and fin in texto:
            idx_ini = texto.find(inicio)
            idx_fin = texto.find(fin) + len(fin)
            frase_a_borrar = texto[idx_ini:idx_fin]
            reemplazar_en_parrafo(para, frase_a_borrar, "")

def limpiar_marcadores_residuales(doc):
    """Remueve cualquier marcador delimitador remanente en el documento."""
    marcadores = [
        "{{INICIA_ATT}}", "{{TERMINA_ATT}}",
        "{{INICIA_TELCEL}}", "{{TERMINA_TELCEL}}",
        "{{INICIA_MOVISTAR}}", "{{TERMINA_MOVISTAR}}",
        "{{INICIA_BAIT}}", "{{TERMINA_BAIT}}",
        "{{INICIA_HECHO7}}", "{{TERMINA_HECHO7}}",
        "{{INICIA_CONCEPTO_SEXTO}}", "{{TERMINA_CONCEPTO_SEXTO}}",
        "{{INICIA_SUPLENCIA}}", "{{TERMINA_SUPLENCIA}}",
        "{{INICIA_RELACION_CONCEPTO_SEXTO}}", "{{TERMINA_RELACION_CONCEPTO_SEXTO}}",
        "{{INICIA_EN_LINEA}}", "{{TERMINA_EN_LINEA}}",
    ]
    for m in marcadores:
        reemplazar(doc, m + " ", "")
        reemplazar(doc, " " + m, "")
        reemplazar(doc, m, "")

def limpiar_espacios_y_sangrias(doc):
    """Depura dobles espacios generados por la interpolación de cadenas."""
    for para in obtener_todos_los_parrafos(doc):
        if para.text:
            if para.runs and para.runs[0].text.startswith(" ") and not para.text.strip().startswith(("•", "-", "1.", "2.", "3.", "4.", "5.")):
                para.runs[0].text = para.runs[0].text.lstrip(" ")
            for run in para.runs:
                if "  " in run.text:
                    run.text = run.text.replace("  ", " ")

def reconstruir_parrafo_con_negritas(para, texto_completo, frases_bold):
    """
    Reconstruye el párrafo aplicando negrita únicamente a las frases de la lista,
    preservando el tipo y tamaño de fuente original.
    """
    fuente = para.runs[0].font.name if para.runs else None
    tamano = para.runs[0].font.size if para.runs else None
    
    frases_validas = [f for f in set(frases_bold) if f and f in texto_completo]
    frases_validas.sort(key=len, reverse=True)
    
    intervalos = []
    for f in frases_validas:
        start = 0
        while True:
            idx = texto_completo.find(f, start)
            if idx == -1:
                break
            fin = idx + len(f)
            if not any(max(idx, b_ini) < min(fin, b_fin) for b_ini, b_fin in intervalos):
                intervalos.append((idx, fin))
            start = idx + 1
            
    intervalos.sort(key=lambda x: x[0])
    
    segmentos = []
    pos_actual = 0
    for b_ini, b_fin in intervalos:
        if b_ini > pos_actual:
            segmentos.append((texto_completo[pos_actual:b_ini], False))
        segmentos.append((texto_completo[b_ini:b_fin], True))
        pos_actual = b_fin
    if pos_actual < len(texto_completo):
        segmentos.append((texto_completo[pos_actual:], False))
        
    para.text = ""
    for texto_seg, es_bold in segmentos:
        run = para.add_run(texto_seg)
        run.bold = es_bold
        if fuente:
            run.font.name = fuente
        if tamano:
            run.font.size = tamano

def tiene_nota_al_pie(run):
    """Verifica si un run contiene el nodo XML de una nota al pie de página."""
    return len(run._r.xpath('.//w:footnoteReference | .//w:footnoteRef')) > 0

def aplicar_formato_proemio_seguro_con_notas(doc, campos_bold_proemio):
    """Aplica negritas en el proemio respetando intactos los nodos XML de notas al pie y citas."""
    for para in obtener_todos_los_parrafos(doc):
        texto_strip = para.text.strip()
        if not texto_strip:
            continue

        # 1. Rubros (PERSONA QUEJOSA, ASUNTO) en negrita
        if texto_strip.startswith('PERSONA QUEJOSA:') or texto_strip.startswith('ASUNTO:'):
            for run in para.runs:
                run.bold = True
            continue

        # 2. Autoridad Jurisdiccional y encabezados en negrita
        if (texto_strip.startswith('C. JUEZ') or texto_strip.startswith('C. JUEZA')
            or texto_strip in ['PRESENTE:', 'EXPONER:'] or texto_strip.startswith('EXPONER:')):
            for run in para.runs:
                run.bold = True
            continue

        # 3. Párrafo de comparecencia (Proemio)
        if 'comparezco para,' in texto_strip or 'señalando como domicilio' in texto_strip:
            for run in para.runs:
                if tiene_nota_al_pie(run):
                    continue

                texto_run = run.text.strip()
                if not texto_run:
                    continue

                es_campo_bold = any(
                    (campo.strip() == texto_run or (len(texto_run) > 3 and texto_run in campo)
                     or (len(campo) > 3 and campo in texto_run))
                    for campo in campos_bold_proemio if campo
                )

                if 'C.P.' in texto_run:
                    es_campo_bold = True

                if any(frase in texto_run for frase in [
                    'Acuerdo General', 'Ley de Amparo', 'artículo', 'LAmp',
                    'comparezco para,', 'señalando como medio', 'términos más amplios',
                ]):
                    es_campo_bold = False

                run.bold = es_campo_bold

def renumerar_y_formatear_pruebas(doc, campos_bold_pruebas):
    """
    Renumera secuencialmente las pruebas, pone el título en negrita y
    resalta términos clave internos.
    """
    contador = 1
    parrafos_a_eliminar = []
    
    en_capitulo_pruebas = False
    for para in obtener_todos_los_parrafos(doc):
        texto = para.text.strip()
        
        if "PRUEBAS:" in texto or texto == "PRUEBAS":
            en_capitulo_pruebas = True
            for r in para.runs:
                r.bold = True
            continue
            
        if en_capitulo_pruebas and ("PIDO:" in texto or "PROTESTO LO NECESARIO" in texto):
            en_capitulo_pruebas = False
            break
            
        if en_capitulo_pruebas and not texto:
            parrafos_a_eliminar.append(para)

    for para in parrafos_a_eliminar:
        p = para._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    for para in obtener_todos_los_parrafos(doc):
        texto = para.text.strip()
        
        es_prueba = (
            texto.startswith("DOCUMENTAL PRIVADA") or
            texto.startswith("PRESUNCIONAL LEGAL") or
            texto.startswith("INSTRUMENTAL DE ACTUACIONES") or
            (len(texto) > 2 and texto[0].isdigit() and texto[1] in ". )" and
             ("DOCUMENTAL" in texto or "PRESUNCIONAL" in texto or "INSTRUMENTAL" in texto))
        )
        
        if es_prueba:
            if texto[0].isdigit() and texto[1] in ". )":
                partes = texto.split(".", 1) if "." in texto[:4] else texto.split(" ", 1)
                texto_sin_num = partes[-1].strip()
            else:
                texto_sin_num = texto

            if "," in texto_sin_num:
                titulo_prueba, desc_prueba = texto_sin_num.split(",", 1)
                titulo_final = f"{contador}. {titulo_prueba},"
                texto_unificado = f"{titulo_final} {desc_prueba.strip()}"
            else:
                titulo_final = f"{contador}. {texto_sin_num}"
                texto_unificado = titulo_final

            frases_bold_esta_prueba = [titulo_final] + campos_bold_pruebas
            reconstruir_parrafo_con_negritas(para, texto_unificado, frases_bold_esta_prueba)
            
            contador += 1

def limpiar_nombre_archivo(texto):
    """
    Elimina acentos, espacios, guiones, símbolos y caracteres especiales.
    Genera una cadena estrictamente alfanumérica para el Portal del PJF.
    """
    if not texto:
        return ''
    texto_sin_acentos = ''.join(
        c for c in unicodedata.normalize('NFD', str(texto))
        if unicodedata.category(c) != 'Mn'
    )
    texto_sin_acentos = (
        texto_sin_acentos.replace('&', 'Y')
        .replace('+', 'Mas')
        .replace('/', '')
        .replace('\\', '')
    )
    nombre_limpio = re.sub(r'[^a-zA-Z0-9]', '', texto_sin_acentos)
    return nombre_limpio

# ============================================================
# 3. INTERFAZ DE USUARIO (STREAMLIT)
# ============================================================
st.set_page_config(page_title="GENERADOR DE DEMANDA DE AMPARO INDIRECTO", layout="wide")

st.title("GENERADOR DE DEMANDA DE AMPARO INDIRECTO")
st.caption("LMTR + Lineamientos de Identificación de Líneas Telefónicas Móviles")

st.warning("""
**AVISO IMPORTANTE**  
Esta es una herramienta de apoyo, **no sustituye la asesoría profesional de un abogado**.  
El documento generado debe ser **revisado y validado** por un profesional del derecho antes de presentarse.
""")

tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
    "1. Datos del quejoso",
    "2. Línea y notificación",
    "3. Abogados",
    "4. Concesionaria",
    "5. Opciones especiales",
    "6. Generar",
    "7. Guía y FAQs"
])

# --- PESTAÑA 1: DATOS DEL QUEJOSO ---
with tab1:
    col_d1, col_d2 = st.columns([2, 1])
    with col_d1:
        justiciable = st.text_input("Nombre completo de la persona quejosa (o del menor) *")
        domicilio = st.text_input("Domicilio (Calle, Número exterior/interior, Colonia) *")
    with col_d2:
        cp = st.text_input("Código Postal (C.P.) *", value="20000")
        ciudad_edo = st.text_input("Ciudad y Estado (Para fecha y firma) *", value="Aguascalientes, Ags.")
    
    st.subheader("Edad y Representación Legal")
    edad = st.radio("Situación de edad *", ["Mayor de edad", "Menor de edad"], horizontal=True)
    
    nombre_representante = ""
    modalidad_menor = ""
    
    if edad == "Menor de edad":
        modalidad_menor = st.radio(
            "Modalidad de representación procesal *",
            [
                "Por su propio derecho (Autónoma - Art. 8º Ley de Amparo)",
                "Por conducto de Padre / Madre / Tutor (Patria Potestad o Tutela)"
            ],
            help="Elige 'Por conducto...' si la demanda se presentará en línea con la FIREL/e.firma del adulto."
        )
        
        if "Padre / Madre / Tutor" in modalidad_menor:
            nombre_representante = st.text_input("Nombre completo del Padre, Madre o Tutor que promueve *")
            
        nombre_promovente = justiciable if justiciable else "[Nombre del promovente]"
        st.warning(f"""
        ### ⚠️ ALERTA - PROMOVENTE MENOR DE EDAD (ART. 8º LEY DE AMPARO)
        Has seleccionado generar una demanda para una persona menor de edad (**{nombre_promovente}**).

        **Ten en cuenta lo siguiente:**
        * **Si promueve por su propio derecho (Art. 8º LAmp):** Al admitir la demanda, el Juez de Distrito requerirá la comparecencia de los padres/tutores o le nombrará un Representante Especial (IFDP).
        * **Si promueve por conducto de padre/tutor:** El padre/tutor ingresará la demanda con su usuario y firma electrónica o mediante su firma autógrafa.
        """)

    st.subheader("Forma de presentación")
    presenta_en_linea = st.checkbox(
        "¿Se presenta en línea (Portal de Servicios en Línea del PJF)?",
        value=True,
        help="Si se marca, se solicitará el usuario del Portal y se incluirá la firma electrónica."
    )
    
    usuario_pjf = ""
    if presenta_en_linea:
        usuario_pjf = st.text_input("Usuario del Portal de Servicios en Línea del PJF *")
        st.caption("Se conservará el bloque de firma electrónica.")
    else:
        st.caption("Se eliminará el bloque de firma electrónica (presentación tradicional).")

# --- PESTAÑA 2: LÍNEA Y NOTIFICACIÓN ---
with tab2:
    numero_telefono = st.text_input("Número de teléfono suspendido *", value="4491234567")
    
    fecha_base = obtener_fecha_base_calendario(numero_telefono)
    fecha_limite_72h = fecha_base + timedelta(days=3)
    digito_detectado = [c for c in numero_telefono if c.isdigit()][-1] if any(c.isdigit() for c in numero_telefono) else "0"
    
    st.info(f"""
    💡 **Ventana de suspensión (Transitorio SEGUNDO - Acuerdo de modificación a Lineamientos):** Para el dígito **{digito_detectado}**, 
    la fecha del calendario es el **{fecha_base.strftime('%d/%m/%Y')}**. La suspensión puede ocurrir 
    en cualquier momento dentro del plazo máximo de 72 horas (hasta el **{fecha_limite_72h.strftime('%d/%m/%Y')}**).
    """)

    col_med, col_fec, col_hor = st.columns([2, 2, 1])
    with col_med:
        medio_notif = st.selectbox(
            "Medio de notificación *",
            ["correo electrónico", "SMS", "correo electrónico/SMS", "aplicación de la concesionaria"]
        )
    with col_fec:
        fecha_notif = st.date_input(
            "Fecha del aviso o suspensión *", 
            value=fecha_limite_72h
        )
    with col_hor:
        hora_notif = st.time_input(
            "Hora del aviso *", 
            value=st.session_state.hora_defecto,
            key="hora_notif_input"
        )
    
    if fecha_notif:
        plazos = calcular_plazos(fecha_notif)
        hora_formato = hora_notif.strftime("%H:%M")
        st.success(f"""
**Datos del Acto Reclamado registrados:**  
• Momento del corte / notificación: **{plazos['str_notif']} a las {hora_formato} hrs** vía {medio_notif}  
• Surte efectos: **{plazos['str_surte']}**  
• Inicia plazo de 15 días hábiles: **{plazos['str_empieza']}**  
• Vencimiento legal: **{plazos['str_final']}**
""")

# --- PESTAÑA 3: ABOGADOS ---
with tab3:
    st.subheader("Autorización de abogados")
    autorizar_abogados = st.checkbox("¿Deseas autorizar abogados?", value=True)
    
    abogado_1 = cedula_1 = usuario_abogado_1 = abogado_2 = cedula_2 = usuario_abogado_2 = ""
    if autorizar_abogados:
        st.warning("""
        🚨 **IMPORTANTE:**  
        Si autorizas a uno, dos o más profesionales del Derecho (artículo 12 de la Ley de Amparo), **es necesario que estén enterados, conformes y en disposición de dar seguimiento al expediente**.  
        
        *Señalar a un abogado sin su conocimiento o consentimiento no servirá*, ya que no revisará los acuerdos ni desahogará prevenciones o recursos a tiempo, lo que puede provocar el desechamiento o la pérdida indeseada del juicio.
        """)
        
        col_ab1, col_ab2 = st.columns(2)
        with col_ab1:
            st.markdown("#### Abogado 1 (Obligatorio)")
            abogado_1 = st.text_input("Nombre del abogado 1 *")
            cedula_1 = st.text_input("Cédula profesional abogado 1 *")
            usuario_abogado_1 = st.text_input("Usuario - Portal de Servicios en Línea del PJF abogado 1 (opcional)")
            
        with col_ab2:
            st.markdown("#### Abogado 2 (Opcional)")
            abogado_2 = st.text_input("Nombre del abogado 2")
            cedula_2 = st.text_input("Cédula profesional abogado 2")
            usuario_abogado_2 = st.text_input("Usuario - Portal de Servicios en Línea del PJF abogado 2 (opcional)")
    else:
        st.caption("Se eliminará la autorización de abogados del proemio.")

# --- PESTAÑA 4: CONCESIONARIA ---
with tab4:
    st.subheader("Selección de Concesionaria")

    carrier = st.selectbox(
        "Concesionaria responsable *",
        ["AT&T", "Telcel", "Movistar", "Bait"],
        help="Selecciona la compañía proveedora de tu servicio móvil.",
    )

    st.info("""
    📋 **Avisos de Privacidad Vigentes y Verificados:**
    * **AT&T:** 04 de junio de 2026
    * **Telcel:** Junio 2025
    * **Movistar:** 15 de marzo de 2026
    * **Bait:** 27 de mayo de 2026
    
    *Última revisión y cotejo de vigencia: **14 de agosto de 2026**.*
    """)

# --- PESTAÑA 5: OPCIONES ESPECIALES ---
with tab5:
    st.subheader("Cuenta bancaria (opcional)")
    tiene_cuenta = st.checkbox("La línea telefónica está vinculada a una cuenta bancaria", value=False)
    
    banco = tipo_cuenta = ""
    if tiene_cuenta:
        banco = st.text_input("Nombre del banco *")
        tipo_cuenta = st.text_input("Tipo / nivel de cuenta (ej. nivel 4) *")
        st.markdown("""
        Nota: La denominación del banco puede consultarse en la [CNBV](https://www.cnbv.gob.mx/Entidades-Autorizadas/Paginas/Banca-Multiple.aspx). El tipo/nivel de cuenta dependerá de las operaciones y en la app del banco suele indicar el tipo/nivel, para mayor referencia revisar en la siguiente nota de ["EL ECONOMISTA"](https://www.eleconomista.com.mx/finanzaspersonales/spei-2-0-cuentas-bancarias-nivel-tienes-correcta-20260623-819730.html).
        """)
    else:
        st.caption("Se eliminará el concepto SEXTO, el hecho relacionado y la prueba bancaria.")

    st.subheader("Suplencia de la Queja")
    
    opciones_vulnerabilidad_menor = [
        "persona con discapacidad",
        "miembro de comunidad o pueblo indígena"
    ]
    
    opciones_vulnerabilidad_mayor = [
        "persona con discapacidad",
        "persona mayor sujeta a estado de interdicción",
        "miembro de comunidad o pueblo indígena",
        "adulto mayor pensionado"
    ]

    if edad == "Menor de edad":
        st.info("Al ser menor de edad, se activa automáticamente la suplencia de la queja (arts. 8 y 79 LAmp).")
        incluir_suplencia = True
        otras = st.multiselect(
            "Otras condiciones de vulnerabilidad compatibles (opcionales):",
            opciones_vulnerabilidad_menor,
            default=[]
        )
        condiciones = ["menor de edad"] + otras
    else:
        incluir_suplencia = st.checkbox("Incluir el capítulo de Suplencia de la Queja", value=False)
        condiciones = []
        if incluir_suplencia:
            condiciones = st.multiselect(
                "Seleccione las condiciones de vulnerabilidad:",
                opciones_vulnerabilidad_mayor,
                default=[]
            )
    
    condicion_texto = ""
    if incluir_suplencia:
        if condiciones:
            condicion_texto = condiciones[0] if len(condiciones) == 1 else ", ".join(condiciones[:-1]) + " y " + condiciones[-1]
        else:
            condicion_texto = "persona en situación de vulnerabilidad"
    else:
        st.caption("No se incluirá el capítulo de Suplencia de la Queja.")

# --- PESTAÑA 6: GENERAR DOCUMENTO ---
with tab6:
    fecha_presentacion = st.date_input(
        "Fecha de presentación *", 
        value=fecha_notif,
        help="Se sugiere por defecto la misma fecha en que conociste el acto reclamado."
    )
    
    st.markdown("""
    Nota: La demanda de amparo puede promoverse desde el mismo día en que se deshabilita/suspende, de conformidad con los criterios jurisprudenciales con registros digitales: [2014596](https://sjf2.scjn.gob.mx/detalle/tesis/2014596), [2010884](https://sjf2.scjn.gob.mx/detalle/tesis/2010884) y [2011123](https://sjf2.scjn.gob.mx/detalle/tesis/2011123).
    """)
    
    st.markdown("---")

    if st.button("Generar documento Word", type="primary", use_container_width=True):
        if not justiciable or not domicilio or not cp or not ciudad_edo or not numero_telefono:
            st.error("Completa los campos obligatorios (*)")
        elif edad == "Menor de edad" and "Padre / Madre / Tutor" in modalidad_menor and not nombre_representante:
            st.error("Si el menor es representado, debes ingresar el nombre del Padre, Madre o Tutor.")
        elif presenta_en_linea and not usuario_pjf:
            st.error("Si se presenta en línea, el usuario del Portal es obligatorio.")
        elif autorizar_abogados and (not abogado_1 or not cedula_1):
            st.error("Si autorizas abogados, el Abogado 1 y su cédula son obligatorios.")
        else:
            with st.spinner("Generando demanda..."):
                try:
                    doc = Document("amparo_template_clean.docx")
                    plazos = calcular_plazos(fecha_notif)

                    # 1. NOTIFICACIÓN Y USUARIO DEL PORTAL
                    if not presenta_en_linea:
                        frase_usuario = (
                            "señalando como medio para oír y recibir todo tipo de notificaciones y documentos, "
                            "aun los de carácter personal, al usuario {{USUARIO}} dentro del Sistema del Portal "
                            "de Servicios en Línea del Poder Judicial de la Federación, usuario que pertenece "
                            "a la suscrita persona quejosa, o en su defecto, por falla técnica, de conformidad "
                            "con el artículo 68 del Acuerdo General 06/2026 del Órgano de Administración de "
                            "Justicia del Poder Judicial de la Federación (en adelante, AG-POAJ-006/2026), "
                            "se realice en el domicilio antes señalado; "
                        )
                        reemplazar_frase(doc, frase_usuario, "")
                    else:
                        if edad == "Menor de edad" and "Padre / Madre / Tutor" in modalidad_menor:
                            reemplazar_frase(
                                doc, 
                                "usuario que pertenece a la suscrita persona quejosa", 
                                "usuario que pertenece a la parte promovente"
                            )
                            reemplazar_frase(
                                doc, 
                                "usuario que pertenece a la persona quejosa", 
                                "usuario que pertenece a la parte promovente"
                            )

                    # 2. AUTORIZACIÓN DE ABOGADOS
                    if not autorizar_abogados:
                        frase_abogados_completa = (
                            "; y autorizando en los términos más amplios del artículo 12 de la Ley de Amparo "
                            "(en adelante, LAmp) al(a los) Lic(s). en Derecho¹ {{ABOGADO_1}} (cédula profesional {{CEDULA_1}}) "
                            "y {{ABOGADO_2}} (cédula profesional {{CEDULA_2}})"
                        )
                        reemplazar_frase(doc, frase_abogados_completa, "")
                        reemplazar_frase(doc, frase_abogados_completa.replace("Derecho¹", "Derecho"), "")
                    elif not (abogado_2 and cedula_2):
                        reemplazar_frase(doc, " y {{ABOGADO_2}} (cédula profesional {{CEDULA_2}})", "")

                    # 3. COMPAÑÍAS TELEFÓNICAS (ELIMINAR NO SELECCIONADAS)
                    markers = {
                        "AT&T": ("{{INICIA_ATT}}", "{{TERMINA_ATT}}"),
                        "Telcel": ("{{INICIA_TELCEL}}", "{{TERMINA_TELCEL}}"),
                        "Movistar": ("{{INICIA_MOVISTAR}}", "{{TERMINA_MOVISTAR}}"),
                        "Bait": ("{{INICIA_BAIT}}", "{{TERMINA_BAIT}}")
                    }
                    
                    for c, (ini, fin) in markers.items():
                        if c != carrier:
                            eliminar_todas_las_secciones(doc, ini, fin)
                    
                    ini_sel, fin_sel = markers[carrier]
                    reemplazar(doc, ini_sel, "")
                    reemplazar(doc, fin_sel, "")

                    # 4. CUENTA BANCARIA / CONCEPTO SEXTO
                    if not tiene_cuenta:
                        eliminar_seccion(doc, "{{INICIA_HECHO7}}", "{{TERMINA_HECHO7}}")
                        eliminar_seccion(doc, "{{INICIA_CONCEPTO_SEXTO}}", "{{TERMINA_CONCEPTO_SEXTO}}")
                        borrar_frase_entre_marcadores(doc, "{{INICIA_RELACION_CONCEPTO_SEXTO}}", "{{TERMINA_RELACION_CONCEPTO_SEXTO}}")

                    # 5. SUPLENCIA DE LA QUEJA
                    if not incluir_suplencia:
                        eliminar_seccion(doc, "{{INICIA_SUPLENCIA}}", "{{TERMINA_SUPLENCIA}}")

                    # 6. PRESENTACIÓN EN LÍNEA
                    if not presenta_en_linea:
                        eliminar_seccion(doc, "{{INICIA_EN_LINEA}}", "{{TERMINA_EN_LINEA}}")

                    # 7. LIMPIEZA DE MARCADORES RESIDUALES
                    limpiar_marcadores_residuales(doc)

                    # 8. CONSTRUCCIÓN DE PERSONERÍA Y FIRMA
                    if edad == "Mayor de edad":
                        texto_personeria = f"{justiciable.upper()}, mayor de edad, por mi propio derecho"
                        texto_firma = justiciable.upper()
                    else:
                        if "Padre / Madre / Tutor" in modalidad_menor:
                            texto_personeria = f"{nombre_representante.upper()}, en representación de mi hijo(a) menor de edad {justiciable.upper()}, en ejercicio de la patria potestad"
                            texto_firma = f"{nombre_representante.upper()}\n\nEn representación de su hijo(a) menor de edad {justiciable.upper()}"
                        else:
                            texto_personeria = f"{justiciable.upper()}, menor de edad, compareciendo por mi propio derecho de conformidad con el artículo 8º de la Ley de Amparo"
                            texto_firma = justiciable.upper()

                    # Formateo dinámico de Cédula + Usuario PJF
                    val_cedula_1 = cedula_1
                    if usuario_abogado_1.strip():
                        val_cedula_1 += f", usuario del Portal de Servicios en Línea del PJF: {usuario_abogado_1.strip()}"
                        
                    val_cedula_2 = cedula_2
                    if usuario_abogado_2.strip():
                        val_cedula_2 += f", usuario del Portal de Servicios en Línea del PJF: {usuario_abogado_2.strip()}"

                    # 9. CONVERSIÓN DE FECHAS Y SUSTITUCIÓN DE DICCIONARIO
                    meses = ["enero","febrero","marzo","abril","mayo","junio",
                             "julio","agosto","septiembre","octubre","noviembre","diciembre"]
                    fecha_pres_str = f"{fecha_presentacion.day} de {meses[fecha_presentacion.month-1]} de {fecha_presentacion.year}"

                    hora_str = hora_notif.strftime("%H:%M")
                    texto_momento_suspension = f"{plazos['str_notif']}, aproximadamente a las {hora_str} horas"

                    ejecutoras = {
                        "AT&T": "AT&T COMUNICACIONES DIGITALES, S. DE R.L. DE C.V. (AT&T)",
                        "Telcel": "RADIOMÓVIL DIPSA, S.A. DE C.V. (Telcel)",
                        "Movistar": "Pegaso PCS, S.A. de C.V. (Movistar)",
                        "Bait": "WAL-MART INNOVACIÓN S. DE R.L. DE C.V. (comercializadora autorizada bajo la marca Bait)"
                    }

                    concesionaria_nombre = ejecutoras[carrier]

                    reemplazos = {
                        "{{PERSONERIA_PROEMIO}}": texto_personeria,
                        "{{FIRMA_NOMBRE}}": texto_firma,
                        "{{JUSTICIABLE}}": justiciable.upper(),
                        "{{DOMICILIO}}": domicilio,
                        "{{CP}}": cp,
                        "{{CIUDAD_EDO}}": ciudad_edo,
                        "{{USUARIO}}": usuario_pjf if presenta_en_linea else "",
                        "{{ABOGADO_1}}": abogado_1 if autorizar_abogados else "",
                        "{{CEDULA_1}}": val_cedula_1 if autorizar_abogados else "",
                        "{{ABOGADO_2}}": abogado_2 if (autorizar_abogados and abogado_2) else "",
                        "{{CEDULA_2}}": val_cedula_2 if (autorizar_abogados and abogado_2) else "",
                        "{{MEDIO_NOTIFICACION}}": medio_notif,
                        "{{FECHA_NOTIF_SUSPENSION}}": texto_momento_suspension,
                        "{{FECHA_SURTE_EFECTOS}}": plazos["str_surte"],
                        "{{FECHA_EMPIEZA_PLAZO}}": plazos["str_empieza"],
                        "{{COMPUTO_FINAL}}": plazos["str_final"],
                        "{{NUMERO_TELEFONO}}": numero_telefono,
                        "{{BANCO}}": banco if tiene_cuenta else "",
                        "{{TIPO_CUENTA}}": tipo_cuenta if tiene_cuenta else "",
                        "{{CONCESIONARIA}}": concesionaria_nombre,
                        "{{FECHA_PRESENTACION}}": fecha_pres_str,
                        "{{CONDICION}}": condicion_texto,
                        "{{CONDICIÓN_VULNERABLE}}": condicion_texto,
                        "{{CONDICION_VULNERABLE}}": condicion_texto,
                    }

                    # Sustitución al pie
                    reemplazar_frase(doc, "Aguascalientes, Ags., a {{FECHA_PRESENTACION}}", f"{ciudad_edo}, a {fecha_pres_str}")
                    reemplazar_frase(doc, "{{CIUDAD_EDO}}, a {{FECHA_PRESENTACION}}", f"{ciudad_edo}, a {fecha_pres_str}")

                    for ph, val in reemplazos.items():
                        reemplazar(doc, ph, val)

                    # 10. APLICACIÓN DE FORMATOS EN NEGRITA
                    campos_bold_proemio = [
                        texto_personeria,
                        domicilio,
                        f'C.P. {cp}',
                        cp,
                        ciudad_edo,
                    ]
                    if presenta_en_linea and usuario_pjf:
                        campos_bold_proemio.append(usuario_pjf)
                    if autorizar_abogados:
                        if abogado_1:
                            campos_bold_proemio.append(abogado_1)
                        if val_cedula_1:
                            campos_bold_proemio.append(val_cedula_1)
                        if abogado_2:
                            campos_bold_proemio.append(abogado_2)
                        if val_cedula_2:
                            campos_bold_proemio.append(val_cedula_2)

                    campos_bold_pruebas = [
                        concesionaria_nombre,
                        numero_telefono,
                    ]
                    if tiene_cuenta:
                        if banco:
                            campos_bold_pruebas.append(banco)
                        if tipo_cuenta:
                            campos_bold_pruebas.append(tipo_cuenta)

                    # Ejecutar formateadores
                    aplicar_formato_proemio_seguro_con_notas(doc, campos_bold_proemio)
                    renumerar_y_formatear_pruebas(doc, campos_bold_pruebas)
                    limpiar_espacios_y_sangrias(doc)

                    # 11. DESCARGA CON NOMBRE LIMPIO (SIN GUIONES NI CARACTERES ESPECIALES)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    carrier_clean = limpiar_nombre_archivo(carrier)
                    justiciable_clean = limpiar_nombre_archivo(justiciable)[:25]
                    fecha_clean = fecha_presentacion.strftime('%Y%m%d')

                    nombre_archivo_seguro = f'Amparo{carrier_clean}{justiciable_clean}{fecha_clean}.docx'

                    st.success('¡Demanda generada y optimizada exitosamente!')
                    st.download_button(
                        '⬇️ Descargar demanda (.docx)',
                        data=buffer,
                        file_name=nombre_archivo_seguro,
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True,
                    )

                except Exception as e:
                    st.error(f"Error al procesar el archivo: {e}")
                    st.info("Verifica que 'amparo_template_clean.docx' esté presente en la misma carpeta.")

# --- PESTAÑA 7: GUÍA RÁPIDA Y FAQS ---
with tab7:
    st.header("📘 GUÍA RÁPIDA DE PRESENTACIÓN")
    
    st.info("""
    💡 **¿Cuentas con e.firma (SAT) o necesitas tramitar la FIREL (PJF)?**
    * **Si ya cuentas con e.firma vigente del SAT:** **No es necesario tramitar la FIREL**. El Portal de Servicios en Línea del PJF acepta directamente los archivos de tu e.firma (`.cer`, `.key` y contraseña) para iniciar sesión, firmar y presentar la demanda.
    * **Si no cuentas con ninguna y deseas presentar en línea:** Se recomienda tramitar la **FIREL**, ya que obtener la e.firma ante el SAT requiere citas presenciales que suelen demorar semanas, mientras que la FIREL se obtiene de forma ágil y **100% en línea** desde tu equipo y celular.
    """)
    
    st.subheader("1. Tutorial de creación de la FIREL (Firma Electrónica del PJF)")
    st.caption(
        "De conformidad con el [Manual de usuario - Proceso de solicitud de Certificado Digital]"
        "(https://www.firel.pjf.gob.mx/Manual%20Solicitud%20Requerimiento%20y%20App%20FiREL.pdf)."
    )
    
    col_f1, col_f2 = st.columns(2)
    with col_f1:
        st.markdown("""
        **Paso A: Formulación de la Solicitud en la Web**
        1. Ingresa al portal oficial: [https://www.firel.pjf.gob.mx/](https://www.firel.pjf.gob.mx/).
        2. Haz clic en **"Solicitar"** en el apartado *Solicitud de un certificado digital de firma electrónica*.
        3. Lee detalladamente y haz clic en **"Aceptar"** en los Términos y Condiciones de Uso.
        4. **Verificación de correo electrónico:** Captura un correo electrónico válido y haz clic en "Verificar". Llegará un código de confirmación de 5 dígitos (enviado por la UNCOCEFI). Captúralo en el portal y haz clic en "Verificar".
        5. **Captura de CURP:** Ingresa tu CURP. El sistema se interconectará con RENAPO para recuperar automáticamente tus datos. Verifica que sean correctos y selecciona *Formular Solicitud de certificado digital*.
        
        **Paso B: Generación de Llaves y Contraseña**
        1. **Contraseña de la Llave Privada:** Genera una contraseña segura (mínimo 8 caracteres, incluyendo mayúsculas, minúsculas, números y caracteres especiales).
        2. **Descarga de Llave Privada (.key):** El sistema generará tu par de llaves. Haz clic en "Descargar llave privada". Guárdala muy bien en tu computadora; si la pierdes, tendrás que revocar el trámite y comenzar desde cero.
        3. **Validación de Identificación (INE):** Captura la clave CIC de tu credencial para votar. El portal validará en tiempo real ante el INE que tu identificación esté vigente.
        """)
    
    with col_f2:
        st.markdown("""
        **Paso C: Elección de la Modalidad de Certificación**
        * **Vía 1: En Línea (Recomendada - Sin salir de casa):** Selecciona "En línea". El portal arrojará un acuse con un código QR. Descarga la app *FIREL en Línea* en tu celular (iOS/Android), escanea el código QR y realiza la validación biométrica/video en vivo. Tu firma quedará autorizada digitalmente en las siguientes horas.
        * **Vía 2: Presencial (Si no cuentas con smartphone compatible):** Selecciona "Presencial", adjunta en formato digital un comprobante de domicilio (no mayor a 3 meses), agenda cita en el módulo del PJF más cercano y acude con tu INE original, comprobante de domicilio y acta de nacimiento.
        
        **Paso D: Generación del Archivo PFX (.pfx)**
        1. Una vez aprobado tu certificado, regresa a [https://www.firel.pjf.gob.mx/](https://www.firel.pjf.gob.mx/) -> *Solicitar* -> *Generar Archivo PFX (.pfx)*.
        2. Digita tu CURP y haz clic en "Continuar".
        3. Selecciona el archivo `.key` descargado en el Paso B, introduce tu contraseña y haz clic en *Generar PFX*.
        4. Descarga tu archivo `.pfx` definitivo listo para firmar y subir tu amparo en el portal de servicios en línea del PJF.
        """)

    st.markdown("---")
    st.subheader("2. Tutorial de Presentación en Línea del Amparo (Portal de Servicios en Línea del PJF)")
    st.caption("Guía paso a paso sincronizada con las 4 etapas del sistema de presentación de demandas del PJF.")

    st.markdown("""
    **Paso 1: Registro, Acceso e Inicio de Sesión**
    1. **Creación de usuario nuevo (si aún no tienes cuenta):** Ingresa al [Portal de Registro de Usuarios del PJF](https://www.serviciosenlinea.pjf.gob.mx/juicioenlinea/juicioenlinea/Usuario/Registro), captura tus datos personales (nombre(s), primer apellido, segundo apellido, usuario deseado, CURP, correo electrónico y contraseña) y confirma tu cuenta mediante el enlace que recibirás por correo.
    2. **Acceso al portal oficial:** Dirígete a [https://www.serviciosenlinea.pjf.gob.mx/](https://www.serviciosenlinea.pjf.gob.mx/).
    3. **Autenticación:** Haz clic en la pestaña **"Ingresa al portal"** -> **"Juzgados de Distrito y Tribunales de Circuito"** -> captura tu Usuario y Contraseña (o entra con tu certificado digital) -> Selecciona el perfil **"Persona física"**.

    **Paso 2: Información de promovente y Ubicación de Oficina (Pestañas 1 y 2 del Portal)**
    1. Dirígete al menú **"Presenta tu demanda, solicitud o escrito inicial"**.
    2. Completa los datos requeridos del promovente y en la selección de Oficina de Correspondencia elige Aguascalientes, conforme al capítulo de competencia.
       * **Selecciona:** `Oficina de Correspondencia Común de Juzgados de Distrito en el Estado de Aguascalientes con residencia en Aguascalientes`.

    **Paso 3: Ingresa tu documento (Pestaña 3 del Portal)**
    1. **Demanda principal:** Adjunta el archivo `.docx` generado por esta aplicación.
       * *Importante:* Verifica que el nombre del archivo no contenga espacios en blanco al final, comas, acentos ni caracteres especiales (ej. `AmparoMovistarJUAN20260819.docx`).
    2. **Pruebas y Anexos (en formato PDF):** Carga los documentos de soporte:
       * **Si representas a una persona menor de edad:** Acta de nacimiento o documento legal que acredite la patria potestad o tutela.
       * **Pruebas documentales:** Capturas de pantalla del corte/suspensión, contrato o recibo de telefonía, y en su caso, estados de cuenta bancarios vinculados.

    **Paso 4: Pestaña "4. Resumen", Firma Electrónica y Resguardo de Constancias**
    1. **Revisión:** Avanza a la pestaña **"4. Resumen"** para verificar los archivos cargados en la lista.
    2. **Habilitación de firma:** Marca la casilla **`[✔] ¿Deseas firmar los documentos?`**.
    3. **Selección de método de firma:**
       * **Pestaña PFX (FIREL / FIRMA.JUDICIAL):** Selecciona tu archivo `.pfx`, ingresa tu contraseña y haz clic en el botón **"Confirmar firma"**.
       * **Pestaña KEY - CER (e.firma del SAT / FIREL):** Si usas la e.firma del SAT, selecciona tus archivos `.cer` y `.key`, escribe tu contraseña privada y haz clic en **"Confirmar firma"**.
    4. **Envío formal:** Marca la casilla **`[✔] He leído el Aviso de Privacidad y otorgo mi consentimiento...`** y haz clic en el botón **"Presentar"**.
    5. **Descarga obligatoria de constancias (¡Muy importante!):**
       * **Descarga el acuse de recibo electrónico:** Contiene el número de registro, sello digital de tiempo y código QR oficial del PJF.
       * **Descarga el documento firmado electrónicamente:** Guarda la copia del escrito con la cadena digital de firma incorporada por el sistema.
       * *Conserva ambos archivos en tu computadora para cualquier aclaración o seguimiento procesal.*
    """)
    
    st.markdown("---")
    st.subheader("3. Tutorial de presentación presencial (física) del amparo")
    st.write("Si optas por no usar la e.firma o FIREL, debes acudir a la Oficialía de Partes Común de los Juzgados de Distrito con los siguientes **10 juegos completos impresos** (demanda y anexos):")
    
    st.markdown("""
    * **Un (1) Original:** Firmado de puño y letra (firma autógrafa) en la última hoja.
    * **Un (1) Acuse de Recibo:** Copia simple donde la Oficialía de Partes estampará su sello fechador físico.
    * **Cinco (5) Copias de Traslado para Autoridades Responsables:** Una copia simple por cada autoridad señalada (Congreso de la Unión -Diputados y Senadores-, CRT, Presidente de la República y Concesionaria).
    * **Una (1) Copia de Traslado para el Ministerio Público Federal (MP):** Requerida por la Ley de Amparo.
    * **Dos (2) Copias para el Incidente de Suspensión:** Para que el Juez abra de inmediato el incidente por duplicado que evite la suspensión de la línea.
    """)

    st.markdown("---")
    st.subheader("4. En caso de menores de edad")
    st.warning("""
    Ni el SAT (para la e.firma) ni el PJF (para la FIREL) permiten que personas menores de edad tramiten firmas electrónicas por sí mismas, dado que carecen de credencial para votar (INE) vigente y capacidad de ejercicio plena.
    
    * **Ruta A: Presentación en línea (forzosamente por sus padres/tutores):** El documento deberá ser promovido por conducto de quienes ejerzan la patria potestad o tutela. Los padres o tutores firmarán la demanda electrónicamente utilizando su propia FIREL o e.firma.
    * **Ruta B: Presentación presencial autónoma (Artículo 8º de la Ley de Amparo):** Si el menor desea interponer el juicio por sí mismo (incluso sin intervención de sus padres o ante su ausencia), el artículo 8º de la Ley de Amparo lo autoriza expresamente. Deberá acudir físicamente a la Oficialía de Partes con sus 10 juegos impresos y firmar de puño y letra (o huella). El Juez de Distrito dictará las medidas cautelares urgentes y le nombrará de inmediato un Representante Especial (asesor jurídico del Instituto Federal de Defensoría Pública).
    """)

    st.markdown("---")
    st.header("❓ Preguntas Frecuentes (FAQ)")

    with st.expander("1. ¿Este sistema sustituye a un **abogado patrono**?"):
        st.write("""
        No, pero sí democratiza radicalmente el acceso a la defensa constitucional. Este sistema es una herramienta de apoyo diseñada para generar de forma automatizada y rápida plantillas de amparo con un formato preestablecido. El contenido argumentativo está sustentado en conceptos de violación que tuvieron éxito en la anterior integración de la Suprema Corte de Justicia de la Nación (SCJN) y en el propio precedente de invalidez del PANAUT (Acción de Inconstitucionalidad 82/2021 y su acumulada 86/2021).
        
        Al entregarte un documento listo para firmar, te ahorra los costosos honorarios de redacción inicial. No obstante, el amparo es un procedimiento técnico y dinámico, por lo que, tras la presentación, el Juez de Distrito puede emitir acuerdos o requerimientos que exigen respuestas rápidas. Por ello, la plantilla habilita las variables `{{ABOGADO_1}}` y `{{ABOGADO_2}}` (opción de agregar más en el docx) para autorizar a profesionales del derecho autorizados bajo el artículo 12 de la Ley de Amparo, quienes se encargarán del debido seguimiento procesal.
        """)

    with st.expander("2. ¿Por qué se hace de esta manera (colectiva, automatizada y masiva)?"):
        st.write("""
        * **Contrapeso de velocidad:** El Estado y las telefónicas coordinaron un sistema automatizado que desconecta la línea en un plazo fatal de 72 horas. La única manera viable de defendernos ante un ataque sistémico masivo es mediante una defensa ciudadana igualmente automatizada, veloz y escalable.
        * **La prueba de fuego para la soberanía del "nuevo" Poder Judicial:** Tras las reformas constitucionales y la elección de personas juzgadoras por voto popular en 2025, el nuevo Poder Judicial de la Federación ha sostenido públicamente que su mandato emana del pueblo. Si se masifica la presentación de estos amparos y la cantidad de quejosos activos supera el total de votantes que participaron en las elecciones judiciales de 2025 (alrededor de 13 millones de votos), la SCJN se verá en la necesidad de ejercer su facultad de atracción. Esta masificación obligará al nuevo Poder Judicial a demostrar si realmente actúa de forma soberana e independiente.
        """)

    with st.expander("3. ¿Por qué el script y la plantilla solo cubren a 4 concesionarias (Telcel, AT&T, Movistar y Bait) si según la CRT existen 148 en total?"):
        st.write("""
        Se debe a una estrategia de máxima cobertura e impacto de escala. Aunque existen 148 concesionarios y Operadores Móviles Virtuales (OMVs), la inmensa mayoría del mercado nacional está concentrado en estas cuatro marcas. De acuerdo con [datos oficiales de la CRT](https://www.gob.mx/crt/prensa/reporta-crt-144-5-millones-de-lineas-celulares-activas-en-mexico?idiom=es), las 144.5 millones de líneas activas se distribuyen así:
        * **América Móvil (Telcel):** 57.90% del total.
        * **AT&T:** 16.42% del mercado.
        * **Telefónica (Movistar):** 14.05% del mercado.
        * **Grupo Walmart (Bait):** 7.04% del mercado.
        
        Al sumar estas cuatro empresas, cubrimos de golpe el **95.41% de todas las líneas celulares activas en México** (más de 137.8 millones de líneas). El resto de las 144 concesionarias representan conjuntamente el 4.59% del mercado (aproximadamente 6.6 millones de líneas).
        """)

    with st.expander(
        "4. ¿Tengo que entregar mi CURP o identificación al Juez? ¿Mis datos"
        " estarán más seguros en el amparo que en la LMTR?"
    ):
        st.markdown("""
        * **Vía tradicional (Presentación física en ventanilla):**  
          Si presentas tu demanda impresa ante la Oficialía de Partes, únicamente proporcionas tu **nombre, domicilio procesal, número telefónico** y los datos contenidos en los documentos que anexes como pruebas. En esta modalidad **no se te exige ninguna captura biométrica adicional**.

        * **Vía en línea (Portal del PJF y trámite de FIREL):**  
          Si optas por el juicio en línea y no cuentas con *e.firma*, al tramitar la **FIREL** ante el PJF se recabará: clave CURP, identificación oficial (INE), fotografía del rostro, fotografía de las huellas dactilares (del dedo meñique al índice) y tu firma.  
          *Diferencia sustantiva:* Esta información se entrega voluntariamente para ejercer el derecho humano de **acceso a la justicia** y queda bajo la estricta custodia del Poder Judicial de la Federación, protegida por el secreto judicial y la reserva legal, impidiendo su transferencia comercial o administrativa.

        * **El peligro Administrativo de la LMTR (Acceso sin orden judicial):**  
          A diferencia del resguardo judicial, el artículo 183 de la **LMTR** y sus Lineamientos obligan a las telefónicas a entregar la **geolocalización en tiempo real, metadatos de comunicación y registros de los usuarios** a autoridades administrativas y ministeriales mediante un simple oficio, **sin necesidad de una orden judicial previa emitida por un juez**.

        * **Riesgo penal (*Spoofing* y Prisión Preventiva Oficiosa):**  
          Vincular forzosamente tu identidad a la línea celular te expone a riesgos de seguridad. Si un tercero realiza ***spoofing*** (enmascaramiento o clonación del identificador de llamada) para efectuar llamadas de extorsión usando tu número, el padrón te señalará automáticamente como el titular del acto delictivo. Al estar catalogada la extorsión en el **artículo 19 constitucional como delito con Prisión Preventiva Oficiosa**, un usuario inocente corre el riesgo de enfrentar una privación automática de su libertad antes de que la autoridad investigue si la línea fue alterada técnicamente.
        """)

    with st.expander("5. ¿Los datos proporcionados en el script y la plantilla son seguros?"):
        st.write("""
        **Sí, la privacidad es absoluta por diseño técnico:**
        * **Procesamiento 100% Local:** El código corre en la memoria caché local de tu navegador web. Nada es guardado en servidores externos, nubes ni enviado por e-mail.
        * **GitHub y Licencia CC BY-SA 4.0:** Proyecto alojado de forma pública para asegurar la transparencia auditada de su código.
        * **Opción de datos genéricos:** Si sientes desconfianza, puedes rellenar el formulario con datos de prueba (ej. "JUAN PÉREZ" y "5512345678"), generar el `.docx` limpio y modificarlo de forma segura en tu computadora usando cualquier editor de texto (Word, Google Docs, LibreOffice Writer, WPS Writer, ONLYOFFICE y FreeOffice).
        """)

    with st.expander("6. ¿Cómo puedo conseguir la plantilla y el script?"):
        st.write("Ambos recursos se encuentran alojados en un repositorio público de GitHub e interfaces en Netlify con acceso gratuito y universal.")

    with st.expander("7. ¿Qué fin se persigue con esta demanda de amparo?"):
        st.write("""
        1. **Evitar la desconexión:** Que el Juez Federal ordene mantener activa tu línea telefónica y abstenerse de exigirte datos biométricos.
        2. **Garantizar tu patrimonio (Si tu número está vinculado a una cuenta bancaria):** Mantener activa tu línea para que puedas seguir utilizando tus aplicaciones bancarias y tokens móviles.
        3. **Fondo constitucional:** Que el Poder Judicial declare la inconstitucionalidad de la LMTR y lineamientos, al violentarse múltiples Derechos Humanos.
        """)

    st.markdown("---")
    st.header("⚖️ Escenarios de respuesta que podemos esperar de los juzgados")
    st.write("El juicio de amparo es sumamente dinámico e impredecible en sus primeros momentos procesales. Tras la presentación del escrito o una vez que se turna al juzgado, el Juez de Distrito dictará un auto en un plazo de 24 horas en donde comúnmente sucederá alguno de estos escenarios:")
    
    st.markdown("""
    * **Admisión y suspensión provisional:** El escenario idóneo. El Juez admite a trámite la demanda, abre el incidente de suspensión por separado y concede la suspensión provisional para que la telefónica no corte la línea. Una vez que tengas el auto, deberás notificar inmediatamente al operador con la resolución firmada digitalmente por el Juez.
    * **Prevención:** El Juez encuentra alguna imprecisión formal o duda técnica en la demanda. Otorgará un plazo fatal de **5 días hábiles** para responder mediante escrito aclaratorio; de lo contrario, se tendrá por no presentada la demanda.
    * **Desechamiento de plano:** El Juez se declara incompetente por razón de materia y pretende remitir de inmediato el expediente a los juzgados que, a su juicio, cuentan con competencia especializada para conocer del asunto. Por tal motivo, se incorporó un apartado específico en el capítulo de suspensión, con fundamento en el artículo 48 de la Ley de Amparo, a fin de solicitar que, antes de efectuar dicha remisión, el órgano jurisdiccional otorgue la suspensión provisional correspondiente. Lo anterior, en virtud de que existe una obligación legal de pronunciarse sobre esta medida cautelar cuando resulte procedente, garantizando desde el inicio la protección de tu derecho a la conectividad y, en su caso, la preservación de tu patrimonio, evitando que la transferencia de los autos genere una afectación irreparable o de difícil reparación.
    
    ---
    **⚠️ Conclusión Procesal:** Debido a estos escenarios y plazos fatales de días hábiles para responder o interponer recursos de queja inmediatos, es de vital importancia contar con la asesoría y representación de un abogado habilitado en el expediente para dar el debido seguimiento procesal a tu juicio.
    """)

    st.markdown("---")
    st.header("🛠️ Instalación y Licencia")

    col_inst, col_lic = st.columns(2)
    
    with col_inst:
        st.subheader("💻 Instalación local")
        st.code("""
# 1. Clonar el repositorio
git clone https://github.com/jaasaav/amparo-lmtr.git
cd amparo-lmtr

# 2. Instalar dependencias
pip install -r requirements.txt

# 3. Ejecutar la app
streamlit run app.py
        """, language="bash")

    with col_lic:
        st.subheader("📄 Licencia y condiciones de uso")
        st.markdown("""
        Distribuido bajo la Licencia **Creative Commons Atribución-CompartirIgual 4.0 Internacional (CC BY-SA 4.0)**:
        
        * **Gratuidad de la aplicación y plantilla:** Prohibida la venta o cobro directo por la descarga del código o archivos `.docx`.
        * **Honorarios permitidos:** Los abogados pueden cobrar honorarios por asesoría, representación y actos procesales subsecuentes.
        * **Prohibición política:** Prohibido su uso para proselitismo partidista o campañas.
        * **Consulta los términos completos:** [Ver LICENCIA.md en GitHub](https://github.com/jaasaav/amparo-lmtr/blob/main/LICENCIA.md).

        ---
        **⚠️ EXCLUSIÓN DE GARANTÍAS Y RESPONSABILIDAD**
        1. **Sin garantía:** La aplicación se provee "tal cual" sin asegurar resoluciones uniformes en todos los juzgados.
        2. **Revisión obligatoria:** No constituye asesoría automática; requiere validación de un abogado postulante.
        3. **Límite de responsabilidad:** El autor no responde por desechamientos, prevenciones o fallos desfavorables.
        """)

    # --- SECCIÓN DE COLABORACIÓN (PULL REQUESTS) ---
    st.markdown("---")
    st.header("🤝 Convocatoria de Colaboración Abierta y Comunidad")
    st.markdown("""
    Este proyecto es de código abierto. Se invita a **académicos, ONGs, litigantes y programadores** a contribuir con mejoras al código o a la argumentación jurídica mediante **Pull Requests (PR)** en GitHub:
    """)

    col_c1, col_c2 = st.columns(2)
    with col_c1:
        st.markdown("""
        **⚖️ Aportes Jurídicos:**
        * Nuevos conceptos de violación y criterios jurisprudenciales.
        * Mejoras a la plantilla `amparo_template_clean.docx`.
        * Monitoreo de avisos de privacidad de las telefónicas.
        """)
    with col_c2:
        st.markdown("""
        **💻 Aportes Técnicos:**
        * Optimización del procesador XML de Word (`python-docx`).
        * Mejoras en la validación de formularios en Streamlit.
        * Automatización de pruebas de sintaxis de variables.
        """)
        
    st.info("📦 **Repositorio oficial para colaborar:** [github.com/jaasaav/amparo-lmtr](https://github.com/jaasaav/amparo-lmtr)")

# --- PIE DE PÁGINA: CONTACTO Y CRÉDITOS ---
st.markdown("---")

col_c1, col_c2, col_c3 = st.columns(3)
with col_c1:
    st.markdown("🌐 [Red Social X](https://x.com/ToxZeak)")
with col_c2:
    st.markdown("🐙 [Perfil de GitHub](https://github.com/jaasaav)")
with col_c3:
    st.markdown("📦 [Repositorio del Código](https://github.com/jaasaav/amparo-lmtr)")

with st.expander("📚 Adaptación argumentativa"):
    st.markdown("""
    Los conceptos de violación **PRIMERO** al **QUINTO** fueron retomados y estructurados a partir de los argumentos formulados en la **Acción de Inconstitucionalidad interpuesta por el INAI** y el **Amicus Curiae presentado por la Red en Defensa de los Derechos Digitales (R3D)**.
    """)
    
with st.expander("🤖 IA involucrada en la aplicación"):
    st.markdown("""
    Esta aplicación fue diseñada bajo las instrucciones y dirección del usuario **Julio Amador**, en colaboración técnica con modelos de lenguaje (**Grok, Gemini y Gemini Notebook**).
    """)